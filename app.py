from flask import Flask, request, jsonify
from flask_cors import CORS
import struct, zlib, re, base64, subprocess, tempfile, os, sys, io, json, time, uuid
import requests

app = Flask(__name__)
CORS(app)

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

# GROQ_API_KEY endi bu yerda, SERVER MUHIT OZGARUVCHISI sifatida saqlanadi -
# oldin index.html ichida ochiq turgan edi, hamma korishi mumkin edi. Render
# Dashboard > Environment da GROQ_API_KEY nomli ozgaruvchi qoshish kerak.
GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')

# Supabase jadvaliga (edutest_store) endi FAQAT shu server orqali, "service
# role" kaliti bilan kiriladi - bu kalit RLS'ni chetlab otadi va HECH QACHON
# brauzerga chiqarilmaydi. Render Dashboard > Environment'da SUPABASE_URL va
# SUPABASE_SERVICE_KEY nomli ozgaruvchilarni qoshish kerak (Supabase Dashboard
# > Settings > API > "service_role" kaliti - "anon" kaliti EMAS).
SUPABASE_URL = os.environ.get('SUPABASE_URL', 'https://ykpegjtexjwddsfgwwpw.supabase.co')
SUPABASE_SERVICE_KEY = os.environ.get('SUPABASE_SERVICE_KEY', '')
DB_KEY = 'edutest_v3'

# ─── CLOUDFLARE R2 (2026-08-25: asosiy malumot + YANGI rasmlar uchun) ─────
# Supabase bepul tarifi (500MB baza + 1GB storage) kelajakda tor kelishi
# xavotirlanib, asosiy JSON blob VA yangi importlarda yaratiladigan rasmlar
# endi R2'ga (10GB bepul, hech qanday chiqish narxi yoq) kochiriladi.
# MUHIM: eski ~18,456 ta Supabase Storage'dagi rasm KOCHIRILMAYDI - ular
# ozgarishsiz Supabase'da qoladi va ishlashda davom etadi (ularning URL'lari
# allaqachon toliq https:// havola, R2'ga kochirish shart emas edi). Faqat
# BUNDAN KEYINGI yangi rasmlar R2'ga yoziladi.
# Kalitlar HECH QACHON kodda ochiq yozilmaydi - Render Environment'da
# saqlanadi (R2_ACCOUNT_ID, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY,
# R2_ENDPOINT_URL, R2_IMAGES_BUCKET, R2_DATA_BUCKET, R2_PUBLIC_URL).
R2_ACCOUNT_ID = os.environ.get('R2_ACCOUNT_ID', '')
R2_ACCESS_KEY_ID = os.environ.get('R2_ACCESS_KEY_ID', '')
R2_SECRET_ACCESS_KEY = os.environ.get('R2_SECRET_ACCESS_KEY', '')
R2_ENDPOINT_URL = os.environ.get('R2_ENDPOINT_URL', '')
R2_IMAGES_BUCKET = os.environ.get('R2_IMAGES_BUCKET', 'edutest-images')
R2_DATA_BUCKET = os.environ.get('R2_DATA_BUCKET', 'edutest-data')
R2_PUBLIC_URL = os.environ.get('R2_PUBLIC_URL', '').rstrip('/')

_r2_client_cache = None
def _r2_client():
    """R2 uchun S3-mos klient. Agar muhit ozgaruvchilari sozlanmagan bolsa,
    None qaytaradi - chaqiruvchi funksiyalar shu holatda ESKI (Supabase)
    yoiga avtomatik qaytadi, hech narsa buzilmaydi."""
    global _r2_client_cache
    if _r2_client_cache is not None:
        return _r2_client_cache
    if not (R2_ACCOUNT_ID and R2_ACCESS_KEY_ID and R2_SECRET_ACCESS_KEY and R2_ENDPOINT_URL):
        return False
    try:
        import boto3
        from botocore.config import Config
        _r2_client_cache = boto3.client(
            's3',
            endpoint_url=R2_ENDPOINT_URL,
            aws_access_key_id=R2_ACCESS_KEY_ID,
            aws_secret_access_key=R2_SECRET_ACCESS_KEY,
            config=Config(signature_version='s3v4'),
            region_name='auto',
        )
        return _r2_client_cache
    except Exception as e:
        print(f"R2 klient yaratishda xato: {e}", file=sys.stderr)
        _r2_client_cache = False
        return False

def r2_upload_image(img_bytes, ext='png'):
    """PNG/JPEG baytlarni R2 (edutest-images) bucket'iga yuklaydi va PUBLIC
    URL qaytaradi. Agar R2 sozlanmagan bolsa, ESKI usulga (base64 data URL -
    keyin brauzer ozi Supabase Storage'ga yuklaydi) zaxira sifatida qaytadi,
    shunda R2 hali ulanmagan bolsa ham import ishlashda davom etadi."""
    ext = (ext or 'png').lower()
    mime = 'image/jpeg' if ext in ('jpg', 'jpeg') else 'image/png'
    client = _r2_client()
    if not client or not R2_PUBLIC_URL:
        return f'data:{mime};base64,' + base64.b64encode(img_bytes).decode()
    try:
        key = f'{uuid.uuid4().hex}.{"jpg" if ext in ("jpg","jpeg") else "png"}'
        client.put_object(Bucket=R2_IMAGES_BUCKET, Key=key, Body=img_bytes, ContentType=mime)
        return f'{R2_PUBLIC_URL}/{key}'
    except Exception as e:
        print(f"R2 rasm yuklashda xato (base64'ga qaytildi): {e}", file=sys.stderr)
        return f'data:{mime};base64,' + base64.b64encode(img_bytes).decode()

def _r2_data_get(key):
    client = _r2_client()
    if not client:
        return None
    try:
        obj = client.get_object(Bucket=R2_DATA_BUCKET, Key=key)
        return obj['Body'].read()
    except Exception:
        return None

def _r2_data_put(key, data_bytes, content_type='application/json'):
    client = _r2_client()
    if not client:
        raise Exception('R2 sozlanmagan (muhit ozgaruvchilari yoq)')
    client.put_object(Bucket=R2_DATA_BUCKET, Key=key, Body=data_bytes, ContentType=content_type)

def _r2_get_main_db_or_migrate():
    """Asosiy JSON blobni R2'dan oqiydi. Agar R2'da HALI mavjud bolmasa
    (birinchi chaqiriq - R2 yangi ulangan), ESKI Supabase'dagi joriy
    malumotni bir martalik ochirib, R2'ga yozib qoyadi (avtomatik
    migratsiya) - shundan keyin har doim R2'dan oqiladi. Bu foydalanuvchi
    tomonidan qolda SQL yoki migratsiya skripti ishga tushirishni talab
    qilmaydi - butunlay shaffof, birinchi haqiqiy sorov paytida sodir
    boladi."""
    data = _r2_data_get(f'{DB_KEY}.json')
    if data is not None:
        return data
    if not SUPABASE_SERVICE_KEY:
        return None
    try:
        r = requests.get(
            f'{SUPABASE_URL}/rest/v1/edutest_store',
            params={'key': f'eq.{DB_KEY}', 'select': 'value'},
            headers={'apikey': SUPABASE_SERVICE_KEY, 'Authorization': 'Bearer ' + SUPABASE_SERVICE_KEY},
            timeout=15
        )
        rows = r.json() if r.status_code < 400 else []
        if rows and rows[0].get('value'):
            value = rows[0]['value']
            _r2_data_put(f'{DB_KEY}.json', value.encode('utf-8'))
            print("R2 migratsiya: Supabase'dagi asosiy malumot R2'ga muvaffaqiyatli kochirildi", file=sys.stderr)
            return value.encode('utf-8')
    except Exception as e:
        print(f"R2 migratsiya xato (Supabase'dan oqishda): {e}", file=sys.stderr)
    return None

# ─── PUSH BILDIRISHNOMA (Web Push / VAPID) ────────────────────────────────
# Ustozga "yangi natija keldi" yoki "o'quvchi internetsiz imtihon topshirdi,
# internet tiklanganda tekshirish mumkin" kabi xabarlarni brauzer push
# orqali yuborish uchun. VAPID kalitlar Render Dashboard > Environment'da
# saqlanishi kerak (bir marta generatsiya qilingan, doim bir xil qoladi).
VAPID_PUBLIC_KEY = os.environ.get('VAPID_PUBLIC_KEY', '')
VAPID_PRIVATE_KEY = os.environ.get('VAPID_PRIVATE_KEY', '')
VAPID_CLAIM_EMAIL = os.environ.get('VAPID_CLAIM_EMAIL', 'mailto:admin@edutest.local')

@app.route('/db_get', methods=['GET', 'OPTIONS'])
def db_get():
    if request.method == 'OPTIONS':
        return '', 200
    # R2 birinchi ustuvorlik - agar sozlangan bolsa, shundan oqiladi
    # (birinchi chaqiriqda Supabase'dan avtomatik migratsiya qiladi).
    if _r2_client():
        try:
            data = _r2_get_main_db_or_migrate()
            if data is not None:
                return jsonify({'value': data.decode('utf-8')})
            return jsonify({'value': None})
        except Exception as e:
            print(f"db_get (R2) xato: {e}", file=sys.stderr)
            # R2 xatoga uchradi - pastdagi Supabase zaxira yoiga otamiz
    if not SUPABASE_SERVICE_KEY:
        return jsonify({'error': 'Server sozlanmagan (SUPABASE_SERVICE_KEY yoq)'}), 500
    try:
        r = requests.get(
            f'{SUPABASE_URL}/rest/v1/edutest_store',
            params={'key': f'eq.{DB_KEY}', 'select': 'value'},
            headers={
                'apikey': SUPABASE_SERVICE_KEY,
                'Authorization': 'Bearer ' + SUPABASE_SERVICE_KEY,
            },
            timeout=15
        )
        if r.status_code >= 400:
            return jsonify({'error': f'Supabase xato: {r.status_code}'}), 502
        data = r.json()
        if data and len(data) > 0:
            return jsonify({'value': data[0]['value']})
        return jsonify({'value': None})
    except Exception as e:
        print(f"db_get error: {e}", file=sys.stderr)
        return jsonify({'error': str(e)}), 500

@app.route('/db_save', methods=['POST', 'OPTIONS'])
def db_save():
    if request.method == 'OPTIONS':
        return '', 200
    try:
        body = request.get_json(force=True, silent=True) or {}
        value = body.get('value')
        if value is None:
            return jsonify({'error': 'No value'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    # R2 birinchi ustuvorlik
    if _r2_client():
        try:
            _r2_data_put(f'{DB_KEY}.json', value.encode('utf-8'))
            return jsonify({'ok': True})
        except Exception as e:
            print(f"db_save (R2) xato: {e}", file=sys.stderr)
            # R2 xatoga uchradi - pastdagi Supabase zaxira yoiga otamiz
    if not SUPABASE_SERVICE_KEY:
        return jsonify({'error': 'Server sozlanmagan (SUPABASE_SERVICE_KEY yoq)'}), 500
    try:
        r = requests.post(
            f'{SUPABASE_URL}/rest/v1/edutest_store',
            headers={
                'apikey': SUPABASE_SERVICE_KEY,
                'Authorization': 'Bearer ' + SUPABASE_SERVICE_KEY,
                'Content-Type': 'application/json',
                'Prefer': 'resolution=merge-duplicates',
            },
            json={'key': DB_KEY, 'value': value},
            timeout=15
        )
        if r.status_code >= 400:
            return jsonify({'error': f'Supabase xato: {r.status_code} {r.text[:200]}'}), 502
        return jsonify({'ok': True})
    except Exception as e:
        print(f"db_save error: {e}", file=sys.stderr)
        return jsonify({'error': str(e)}), 500

@app.route('/vapid_public_key', methods=['GET'])
def vapid_public_key():
    """Frontend push obunasi yaratishda shu ochiq kalitni so'raydi."""
    if not VAPID_PUBLIC_KEY:
        return jsonify({'error': 'Server sozlanmagan (VAPID_PUBLIC_KEY yoq)'}), 500
    return jsonify({'key': VAPID_PUBLIC_KEY})


@app.route('/notify_teacher', methods=['POST', 'OPTIONS'])
def notify_teacher():
    """Ustozning barcha push obunalariga (bir nechta qurilma bo'lishi
    mumkin) bildirishnoma yuboradi. Obunalar bazadagi (edutest_v3) JSON
    ichida db.pushSubs sifatida saqlanadi - frontend ularni saveDB() orqali
    oddiy ma'lumot kabi Supabase'ga yozadi, bu yerda faqat o'qib, push
    yuboramiz. Eskirgan/bekor qilingan obunalar (410/404 xato qaytarganlar)
    javobda "stale" sifatida qaytariladi - frontend keyingi saveDB'da ularni
    tozalab qo'yadi.
    """
    if request.method == 'OPTIONS':
        return '', 200
    if not VAPID_PRIVATE_KEY or not VAPID_PUBLIC_KEY:
        return jsonify({'error': 'Server sozlanmagan (VAPID kalitlar yoq)'}), 500
    if not SUPABASE_SERVICE_KEY:
        return jsonify({'error': 'Server sozlanmagan (SUPABASE_SERVICE_KEY yoq)'}), 500
    try:
        from pywebpush import webpush, WebPushException
    except ImportError:
        return jsonify({'error': 'pywebpush ornatilmagan'}), 500
    try:
        body = request.get_json(force=True, silent=True) or {}
        teacher_id = body.get('teacherId')
        title = body.get('title', 'EduTest Pro')
        text = body.get('body', '')
        url = body.get('url', '/')
        if not teacher_id:
            return jsonify({'error': 'teacherId kerak'}), 400

        r = requests.get(
            f'{SUPABASE_URL}/rest/v1/edutest_store',
            params={'key': f'eq.{DB_KEY}', 'select': 'value'},
            headers={'apikey': SUPABASE_SERVICE_KEY, 'Authorization': 'Bearer ' + SUPABASE_SERVICE_KEY},
            timeout=15
        )
        if r.status_code >= 400:
            return jsonify({'error': f'Supabase xato: {r.status_code}'}), 502
        rows = r.json()
        if not rows:
            return jsonify({'sent': 0, 'stale': []})
        db = json.loads(rows[0]['value'])
        subs = [s for s in (db.get('pushSubs') or []) if s.get('teacherId') == teacher_id]

        sent, stale = 0, []
        payload = json.dumps({'title': title, 'body': text, 'url': url})
        for sub in subs:
            try:
                webpush(
                    subscription_info=sub.get('subscription'),
                    data=payload,
                    vapid_private_key=VAPID_PRIVATE_KEY,
                    vapid_claims={'sub': VAPID_CLAIM_EMAIL},
                )
                sent += 1
            except WebPushException as e:
                status = getattr(e.response, 'status_code', None)
                if status in (404, 410):
                    stale.append(sub.get('subscription', {}).get('endpoint'))
                print(f"notify_teacher push xato: {e}", file=sys.stderr)
            except Exception as e:
                print(f"notify_teacher push xato: {e}", file=sys.stderr)
        return jsonify({'sent': sent, 'total': len(subs), 'stale': stale})
    except Exception as e:
        print(f"notify_teacher error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


@app.route('/ai_check', methods=['POST', 'OPTIONS'])
def ai_check():
    """Frontend endi Groq'ga togridan-togri emas, shu endpoint orqali murojaat
    qiladi - shunda haqiqiy API kalit brauzerga hech qachon chiqmaydi."""
    if request.method == 'OPTIONS':
        return '', 200
    if not GROQ_API_KEY:
        return jsonify({'error': "AI tekshirish xizmati sozlanmagan (serverda GROQ_API_KEY yoq)"}), 500
    try:
        body = request.get_json(force=True, silent=True) or {}
        messages = body.get('messages')
        if not messages:
            return jsonify({'error': 'No messages'}), 400
        r = requests.post(
            'https://api.groq.com/openai/v1/chat/completions',
            headers={'Content-Type': 'application/json', 'Authorization': 'Bearer ' + GROQ_API_KEY},
            json={
                'model': body.get('model', 'llama-3.3-70b-versatile'),
                'messages': messages,
                'max_tokens': body.get('max_tokens', 1000),
                'temperature': body.get('temperature', 0.1),
            },
            timeout=60
        )
        return jsonify(r.json()), r.status_code
    except Exception as e:
        print(f"ai_check error: {e}", file=sys.stderr)
        return jsonify({'error': str(e)}), 500

def find_xml(data):
    for i in range(len(data)-6, max(0, len(data)-2000000), -1):
        if i+1 >= len(data): continue
        b0, b1 = data[i], data[i+1]
        if b0 == 0x78 and b1 in (0x01, 0x9c, 0xda, 0x5e):
            for method in [
                lambda c: zlib.decompress(c),
                lambda c: zlib.decompress(c, 47),
                lambda c: zlib.decompress(c[2:], -15),
                lambda c: zlib.decompress(c, -15),
            ]:
                try:
                    out = method(data[i:i+1000000])
                    if b'QuestionBlock' in out or b'<?xml' in out:
                        return out.decode('utf-8', errors='replace')
                except: pass
    return None

def convert_all_emfs(emf_list):
    """Barcha EMF larni bitta LibreOffice session da o'girish.
    emf_list: [(idx, emf_bytes), ...] - idx dunyoning istalgan raqami bolishi mumkin"""
    if not emf_list:
        return {}

    tmpdir = tempfile.mkdtemp(prefix='edutest_emf_')
    results = {}  # idx -> base64

    try:
        emf_paths = {}
        for idx, emf_data in emf_list:
            emf_path = os.path.join(tmpdir, f'f{idx}.emf')
            with open(emf_path, 'wb') as f:
                f.write(emf_data)
            emf_paths[idx] = emf_path

        env = os.environ.copy()
        env['HOME'] = tmpdir

        all_emf_paths = list(emf_paths.values())

        print(f"Converting {len(all_emf_paths)} EMFs in batches (1 LO session)...", file=sys.stderr)
        # MUHIM: Render "Free" tarifi atigi 512 MB xotira beradi. Bitta
        # LibreOffice chaqiruvida juda kop fayl bolsa, xotira tugab (OOM)
        # butun server qulab tushishi mumkin edi. Shuning uchun BATCH kichik
        # tutilgan - sekinroq, lekin barqaror.
        BATCH = 15
        for b_start in range(0, len(all_emf_paths), BATCH):
            batch = all_emf_paths[b_start:b_start+BATCH]
            r = subprocess.run(
                ['libreoffice', '--headless', '--norestore',
                 '--convert-to', 'png:draw_png_Export:{PixelWidth:550}',
                 '--outdir', tmpdir] + batch,
                capture_output=True, timeout=300, env=env
            )
            print(f"Batch {b_start//BATCH+1}: rc={r.returncode}", file=sys.stderr)

        for idx, emf_path in emf_paths.items():
            png_path = emf_path.replace('.emf', '.png')
            if os.path.exists(png_path) and os.path.getsize(png_path) > 2000:
                try:
                    from PIL import Image
                    import io
                    img = Image.open(png_path).convert('RGB')
                    bbox = img.point(lambda x: 0 if x > 240 else 255).convert('L').getbbox()
                    if bbox:
                        pad = 15
                        w, h = img.size
                        bbox = (max(0,bbox[0]-pad), max(0,bbox[1]-pad),
                                min(w,bbox[2]+pad), min(h,bbox[3]+pad))
                        img = img.crop(bbox)
                    buf = io.BytesIO()
                    img.save(buf, format='PNG', optimize=True)
                    png_bytes = buf.getvalue()
                    img.close()
                    buf.close()
                except Exception as e:
                    print(f"  crop err: {e}", file=sys.stderr)
                    with open(png_path, 'rb') as f:
                        png_bytes = f.read()
                results[idx] = r2_upload_image(png_bytes, 'png')
                del png_bytes
            else:
                print(f"  EMF[{idx}] -> FAILED", file=sys.stderr)
            # Rasm faylini darhol ochirib, diskni ham bosh qilamiz
            try: os.unlink(png_path)
            except: pass
            try: os.unlink(emf_path)
            except: pass
        import gc
        gc.collect()

    except subprocess.TimeoutExpired:
        print("LO timeout!", file=sys.stderr)
    except Exception as e:
        print(f"LO error: {e}", file=sys.stderr)
    finally:
        for fp in os.listdir(tmpdir):
            try: os.unlink(os.path.join(tmpdir, fp))
            except: pass
        try: os.rmdir(tmpdir)
        except: pass

    return results

def read_rvf(data, pos, length):
    if length <= 0 or pos <= 0:
        return None, None
    rvf = data[pos:pos+length]

    jpg_start = rvf.find(b'\xff\xd8\xff')
    if jpg_start >= 0:
        jpg_data = rvf[jpg_start:]
        end = jpg_data.rfind(b'\xff\xd9')
        if end >= 0: jpg_data = jpg_data[:end+2]
        if len(jpg_data) > 500:
            return None, r2_upload_image(jpg_data, 'jpg')

    png_start = rvf.find(b'\x89PNG\r\n\x1a\n')
    if png_start >= 0:
        png_data = rvf[png_start:]
        if len(png_data) > 500:
            return None, r2_upload_image(png_data, 'png')

    tmet_pos = rvf.find(b'TMetafile\r\n')
    if tmet_pos >= 0:
        after = rvf[tmet_pos+11:]
        # TMetafile'dan keyin "spacing=", "width=", "height=" kabi bir nechta
        # metama'lumot qatorlari kelishi mumkin (obyekt turiga qarab har xil).
        # Hammasini o'tkazib yuboramiz, faqat binary EMF boshlanguncha.
        while True:
            nl = after.find(b'\r\n')
            if nl < 0 or nl > 40:
                break
            line = after[:nl]
            if b'=' in line and all(32 <= c < 127 for c in line):
                after = after[nl+2:]
            else:
                break
        if len(after) >= 8:
            emf_size = struct.unpack_from('<I', after, 0)[0]
            if 100 < emf_size <= len(after) - 4:
                candidate = after[4:4+emf_size]
                if candidate[:4] == b'\x01\x00\x00\x00':
                    return '__EMF__', candidate
            candidate2 = after[4:]
            if len(candidate2) > 100 and candidate2[:4] == b'\x01\x00\x00\x00':
                emf_hdr_size = struct.unpack_from('<I', candidate2, 4)[0]
                if 100 < emf_hdr_size <= len(candidate2):
                    return '__EMF__', candidate2[:emf_hdr_size]
                return '__EMF__', candidate2
            if after[:4] == b'\x01\x00\x00\x00':
                return '__EMF__', after

    lines = rvf.split(b'\r\n')
    if len(lines) >= 3:
        text_part = b'\r\n'.join(lines[2:])
        try:
            text = text_part.decode('utf-16-le', errors='replace').strip()
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text).strip()
            readable = sum(1 for c in text if c.isprintable() or c in '\n\r\t')
            # MUHIM: avval bu chegara 0.7 edi - EMF/binary "chiqindi" baytlar
            # ham UTF-16 sifatida "o'qiladigan" chiqib, xato matn (masalan
            # "a ‰‰币烈鹵鹵▊2") savolga yozilib qolar edi. Endi ancha qattiqroq:
            # deyarli barcha belgilar chop etiladigan (0.95+) VA lotin/raqam/
            # tinish belgilaridan iborat "normal" ulush yetarli (0.85+) bo'lishi
            # shart - aks holda buni matn deb qabul qilmaymiz.
            normal_chars = sum(1 for c in text if c.isalnum() or c in " .,+-*/=()[]{}%'\"?!:;\u2116")
            normal_ratio = normal_chars / max(len(text), 1)
            if len(text) > 2 and readable / max(len(text), 1) > 0.95 and normal_ratio > 0.85:
                return text, None
        except: pass

    # Bu yergacha yetib kelsa - na EMF, na rasm, na ishonchli matn topildi.
    # Chiqindi belgi qo'yib yubormaymiz - "aniqlanmadi" deb ustozga signal
    # beramiz, u keyin "✏️ Tahrirlash" orqali qo'lda to'g'rilaydi.
    return '__UNRECOGNIZED__', None

def extract_questions_raw(xml_text, data):
    """parse_questions bilan bir xil, lekin EMF larni konvertatsiya QILMAYDI -
    faqat xom emf_tasks royxatini qaytaradi. Bu bir nechta faylni birlashtirib,
    hammasini BITTA LibreOffice chaqiruvida ogirish uchun kerak (tezlik uchun)."""
    name_m = re.search(r'<Name>([\s\S]*?)</Name>', xml_text)
    topic = name_m.group(1).strip() if name_m else 'Test'
    qta_m = re.search(r'<QuestionsToAsk>(\d+)</QuestionsToAsk>', xml_text)
    questions_to_ask = int(qta_m.group(1)) if qta_m else 20

    blocks = re.findall(r'<QuestionBlock[^>]*>([\s\S]*?)</QuestionBlock>', xml_text)
    questions = []
    emf_tasks = []  # {'kind':'q'|'a', 'q_idx':i, 'opt_idx':j, 'emf':bytes}

    for i, block in enumerate(blocks):
        type_m = re.search(r'<QuestionTypeName>(.*?)</QuestionTypeName>', block)
        qtype = type_m.group(1).strip() if type_m else 'MultipleChoice'
        if qtype not in ('MultipleChoice', 'MultipleResponse'): continue

        content_m = re.search(r'<Content>([\s\S]*?)</Content>', block)
        q_text = ''
        img_b64 = None
        emf_data_q = None

        if content_m:
            content = content_m.group(1)
            plain_m = re.search(r'<PlainText>([\s\S]*?)</PlainText>', content)
            plain = plain_m.group(1).strip() if plain_m else ''
            rvf_m = re.search(
                r'<RVFStoredPos>(\d+)</RVFStoredPos>\s*<RVFStoredLen>(\d+)</RVFStoredLen>',
                content)
            if rvf_m:
                rp, rl = int(rvf_m.group(1)), int(rvf_m.group(2))
                rt, ri = read_rvf(data, rp, rl)
                if rt == '__EMF__':
                    q_text = plain or '(formula)'
                    emf_data_q = ri
                elif ri:
                    q_text = plain
                    img_b64 = ri
                elif rt == '__UNRECOGNIZED__':
                    q_text = plain or "\u26a0\ufe0f (formula aniqlanmadi - qo'lda kiriting)"
                elif rt:
                    q_text = rt
                else:
                    q_text = plain
            else:
                q_text = plain

        if not q_text and not img_b64 and not emf_data_q:
            continue

        opts, corr = [], []
        ans_emf_tasks = []
        for am in re.finditer(
            r'<Answer\s+IsCorrect="(Yes|No)"[\s\S]*?<Content>([\s\S]*?)</Content>',
            block):
            ac = am.group(2)
            ap = re.search(r'<PlainText>([\s\S]*?)</PlainText>', ac)
            a_plain = ap.group(1).strip() if ap else ''
            a_rvf_m = re.search(
                r'<RVFStoredPos>(\d+)</RVFStoredPos>\s*<RVFStoredLen>(\d+)</RVFStoredLen>',
                ac)
            a_text = a_plain
            a_emf = None
            if a_rvf_m:
                a_pos = int(a_rvf_m.group(1))
                a_len = int(a_rvf_m.group(2))
                a_rt, a_ri = read_rvf(data, a_pos, a_len)
                if a_rt == '__EMF__':
                    a_emf = a_ri
                    a_text = a_plain or '__IMG_PENDING__'
                elif a_ri:
                    a_text = a_ri
                elif a_rt == '__UNRECOGNIZED__':
                    a_text = a_plain or "\u26a0\ufe0f (aniqlanmadi)"
                elif a_rt and len(a_rt) > len(a_plain):
                    a_text = a_rt
            if a_text is not None or a_emf:
                opt_idx = len(opts)
                opts.append(a_text if a_text else '')
                if am.group(1) == 'Yes':
                    corr.append(opt_idx)
                if a_emf:
                    ans_emf_tasks.append((opt_idx, a_emf))

        if len(opts) >= 2 and corr:
            q_obj = {
                'id': i, 'subject': 'math', 'topic': topic,
                'text': q_text or '(Rasm)',
                'options': opts, 'correct': corr,
                'isMulti': (qtype == 'MultipleResponse') or (len(corr) > 1),
            }
            if img_b64:
                q_obj['image'] = img_b64
            q_idx = len(questions)
            questions.append(q_obj)
            if emf_data_q:
                emf_tasks.append({'kind': 'q', 'q_idx': q_idx, 'emf': emf_data_q})
            for opt_idx, a_emf in ans_emf_tasks:
                emf_tasks.append({'kind': 'a', 'q_idx': q_idx, 'opt_idx': opt_idx, 'emf': a_emf})

    return topic, questions, emf_tasks, questions_to_ask

def resolve_emf_tasks(questions, emf_tasks):
    """Bitta fayl uchun: emf_tasks larni konvertatsiya qilib, questions ichiga joylaydi"""
    if not emf_tasks:
        return
    emf_list = [(idx, t['emf']) for idx, t in enumerate(emf_tasks)]
    emf_results = convert_all_emfs(emf_list)
    for idx, b64 in emf_results.items():
        t = emf_tasks[idx]
        if t['kind'] == 'q':
            questions[t['q_idx']]['image'] = b64
        else:
            opts = questions[t['q_idx']]['options']
            if t['opt_idx'] < len(opts):
                opts[t['opt_idx']] = b64

def parse_questions(xml_text, data):
    topic, questions, emf_tasks, questions_to_ask = extract_questions_raw(xml_text, data)
    resolve_emf_tasks(questions, emf_tasks)
    img_count = sum(1 for q in questions if q.get('image'))
    print(f"Done: {len(questions)} questions, {img_count} images", file=sys.stderr)
    return {'topic': topic, 'questions': questions, 'questionsToAsk': questions_to_ask}

@app.route('/parse', methods=['POST', 'OPTIONS'])
def parse():
    if request.method == 'OPTIONS':
        return '', 200
    try:
        content_type = (request.content_type or '')
        if 'application/json' in content_type:
            body = request.get_json(force=True, silent=True) or {}
            b64 = body.get('data', '')
            if not b64:
                return jsonify({'error': 'No data'}), 400
            data = base64.b64decode(b64)
        else:
            data = request.data
        if not data:
            return jsonify({'error': 'No data'}), 400
        print(f"Received: {len(data)} bytes", file=sys.stderr)
        xml_text = find_xml(data)
        if not xml_text:
            return jsonify({'error': 'XML not found'}), 400
        result = parse_questions(xml_text, data)
        return jsonify(result)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


# ─── BOSHQA FAYL TURLARI (Word/PDF/TXT) - AI ORQALI TAHLIL ────────────────
# EasyQuizzy .exe fayllar QATIY strukturaga ega bolgani uchun regex bilan
# ochib olinadi. Lekin Word/PDF/TXT fayllar HAR XIL formatda yozilishi
# mumkin - shuning uchun bu yerda: 1) fayldan matn+rasmlarni (joylashuv
# tartibida, rasm ornida [RASM_N] belgisi bilan) chiqarib olamiz, 2) shu
# matnni AI'ga (Groq) berib, savol/variant/togri-javob strukturasini
# chiqarishni soraymiz, 3) AI qaytargan [RASM_N] belgilarni haqiqiy rasmga
# almashtiramiz.

def extract_docx_content(data):
    """DOCX fayldan paragraflar tartibida matn va rasm ornlarini chiqaradi.
    MUHIM: qalin (bold) va tagiga chizilgan (underline) matnni **belgi**
    bilan saqlab qoladi - aks holda ustoz togri javobni qalin qilib
    belgilagan bolsa ham, bu ma'lumot AI'ga yetib bormay, hech qanday togri
    javob topilmas edi.
    Qaytaradi: (matn, {placeholder: (rasm_bytes, ext)})"""
    import docx
    doc = docx.Document(io.BytesIO(data))
    images = {}
    img_counter = [0]
    ns_a = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    ns_r = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
    lines = []

    def get_image_bytes(rId):
        try:
            part = doc.part.related_parts[rId]
            return part.blob
        except Exception:
            return None

    for para in doc.paragraphs:
        line_parts = []
        for run in para.runs:
            if run.text:
                t = run.text
                is_marked = bool(run.bold) or bool(run.underline) or bool(run.font.highlight_color)
                if is_marked and t.strip():
                    t = '**' + t + '**'
                line_parts.append(t)
            for blip in run._element.findall('.//' + ns_a + 'blip'):
                rId = blip.get(ns_r + 'embed')
                if rId:
                    img_bytes = get_image_bytes(rId)
                    if img_bytes:
                        img_counter[0] += 1
                        ph = f'[RASM_{img_counter[0]}]'
                        images[ph] = (img_bytes, 'png')
                        line_parts.append(ph)
        if line_parts:
            lines.append(''.join(line_parts))
    return '\n'.join(lines), images


def extract_pdf_content(data):
    """PDF fayldan sahifa boyicha matn va rasmlarni Y-koordinata (joylashuv)
    tartibida chiqaradi. Qalin (bold) matn ham **belgi** bilan saqlanadi -
    PyMuPDF span flags orqali aniqlanadi (bit4 = bold)."""
    import fitz
    doc = fitz.open(stream=data, filetype='pdf')
    images = {}
    img_counter = [0]
    lines = []
    for page in doc:
        blocks = page.get_text('dict').get('blocks', [])
        items = []
        for b in blocks:
            if b.get('type') == 0:
                text = ''
                for line in b.get('lines', []):
                    for span in line.get('spans', []):
                        span_text = span.get('text', '')
                        is_bold = bool(span.get('flags', 0) & 16) or ('bold' in span.get('font', '').lower())
                        if is_bold and span_text.strip():
                            span_text = '**' + span_text + '**'
                        text += span_text
                    text += '\n'
                if text.strip():
                    items.append((b['bbox'][1], text.strip()))
            elif b.get('type') == 1:
                img_bytes = b.get('image')
                ext = b.get('ext', 'png')
                if img_bytes:
                    img_counter[0] += 1
                    ph = f'[RASM_{img_counter[0]}]'
                    images[ph] = (img_bytes, ext)
                    items.append((b['bbox'][1], ph))
        items.sort(key=lambda x: x[0])
        for _, content in items:
            lines.append(content)
    doc.close()
    return '\n'.join(lines), images


def extract_content_from_file(filename, data):
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'docx':
        return extract_docx_content(data)
    elif ext == 'pdf':
        return extract_pdf_content(data)
    elif ext == 'txt':
        return data.decode('utf-8', errors='replace'), {}
    return None, {}


def parse_text_with_ai(text):
    if not GROQ_API_KEY:
        raise Exception('AI xizmati sozlanmagan (serverda GROQ_API_KEY yoq)')
    # MUHIM: Groq bepul tarifida 12000 token/daqiqa (TPM) limiti bor -
    # kirish (matn) VA chiqish (max_tokens) yigindisiga taalluqli.
    MAX_CHARS = 20000
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
    prompt = (
        "Quyidagi hujjat matnidan test savollarini JSON korinishida chiqarib ol.\n"
        "QATIY QOIDALAR:\n"
        "- FAQAT JSON qaytar, boshqa hech qanday matn, izoh yoki markdown yozma.\n"
        "- Format: {\"topic\": \"mavzu nomi\", \"questions\": [{\"text\": \"savol matni\", "
        "\"options\": [\"variant1\",\"variant2\"], \"correct\": [0], \"isMulti\": false}]}\n"
        "- \"correct\" - togri javob(lar)ning options ichidagi index(lar)i (0 dan boshlanadi).\n"
        "- Agar bir nechta togri javob bolsa, isMulti:true va correct bir nechta index bolsin.\n"
        "- MUHIM: matnda **ikki yulduzcha orasidagi** qism - bu asl hujjatda QALIN, TAGIGA "
        "CHIZILGAN yoki BELGILANGAN (highlight) matnni bildiradi. O'qituvchilar odatda TOGRI "
        "JAVOBNI aynan shu tarzda (qalin qilib) belgilaydilar - shuning uchun variant matni "
        "**qalin** bolsa, o'sha variant TOGRI JAVOB deb hisobla. Variantning FAQAT bir qismi "
        "qalin bolsa ham (masalan raqami yoki bitta sozi), butun variant togri hisoblanadi.\n"
        "- Agar hech qanday variant qalin qilinmagan VA \"Javob:\", \"togri:\" kabi aniq izoh "
        "ham yoq bolsa, o'sha savolni OTKAZIB YUBOR - taxmin qilma.\n"
        "- JSON javobingda **belgilarini OLIB TASHLA (ular faqat sen uchun ishora, natijaviy "
        "matnda bolmasligi kerak) - masalan \"**4**\" emas, faqat \"4\" deb yoz.\n"
        "- Matnda [RASM_N] kabi belgilar bolishi mumkin - bular rasm ornini bildiradi. Bu "
        "belgilarni ANIQ ozgarishsiz saqlab qol. Agar savol yoki variant BUTUNLAY rasmdan "
        "iborat bolsa, oshu maydonda FAQAT placeholder'ni yoz (masalan faqat \"[RASM_2]\"), "
        "boshqa matn qoshma.\n"
        "- \"topic\" - hujjat sarlavhasi yoki mazmuniga mos qisqa nom.\n\n"
        "Matn:\n" + text
    )
    r = requests.post(
        'https://api.groq.com/openai/v1/chat/completions',
        headers={'Content-Type': 'application/json', 'Authorization': 'Bearer ' + GROQ_API_KEY},
        json={
            'model': 'llama-3.3-70b-versatile',
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': 4000,
            'temperature': 0.1,
            'response_format': {'type': 'json_object'},
        },
        timeout=90,
    )
    result = r.json()
    if 'error' in result:
        raise Exception(str(result['error']))
    content = result['choices'][0]['message']['content'].strip()
    if content.startswith('```'):
        content = re.sub(r'^```[a-zA-Z]*\n?', '', content)
        content = re.sub(r'```\s*$', '', content)
    return json.loads(content)


def substitute_image_placeholders(questions, images):
    def resolve(s):
        if not isinstance(s, str):
            return s
        m = re.fullmatch(r'\[RASM_(\d+)\]', s.strip())
        if not m:
            return s
        ph = f'[RASM_{m.group(1)}]'
        if ph not in images:
            return s
        img_bytes, ext = images[ph]
        mime_ext = 'jpg' if ext.lower() in ('jpg', 'jpeg') else 'png'
        try:
            from PIL import Image
            im = Image.open(io.BytesIO(img_bytes)).convert('RGB')
            buf = io.BytesIO()
            im.save(buf, format='PNG', optimize=True)
            img_bytes = buf.getvalue()
            mime_ext = 'png'
        except Exception:
            pass
        return r2_upload_image(img_bytes, mime_ext)

    for q in questions:
        qtext = (q.get('text') or '').strip()
        resolved = resolve(qtext)
        if resolved != qtext:
            q['img'] = resolved
            q['text'] = ''
        q['options'] = [resolve(o) for o in q.get('options', [])]
    return questions


@app.route('/parse_document', methods=['POST', 'OPTIONS'])
def parse_document():
    """Word (.docx), PDF (.pdf) va oddiy matn (.txt) fayllardan savol
    import qilish - AI (Groq) yordamida erkin formatdagi matnni
    strukturaga solib beradi."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        body = request.get_json(force=True, silent=True) or {}
        b64 = body.get('data', '')
        filename = body.get('filename', 'fayl')
        if not b64:
            return jsonify({'error': 'No data'}), 400
        data = base64.b64decode(b64)
        text, images = extract_content_from_file(filename, data)
        if text is None:
            return jsonify({'error': "Fayl turi qollab-quvvatlanmaydi (.docx, .pdf, .txt qollab-quvvatlanadi)"}), 400
        if not text.strip():
            return jsonify({'error': 'Fayldan matn topilmadi'}), 400
        parsed = parse_text_with_ai(text)
        questions = parsed.get('questions', [])
        if not questions:
            return jsonify({'error': "Fayldan hech qanday savol aniqlanmadi (togri javob belgilanmagan bolishi mumkin)"}), 400
        questions = substitute_image_placeholders(questions, images)
        for i, q in enumerate(questions):
            q['id'] = f'{int(time.time()*1000)}_{i}_{uuid.uuid4().hex[:6]}'
            q.setdefault('subject', 'math')
            if 'isMulti' not in q:
                q['isMulti'] = len(q.get('correct', [])) > 1
        topic = parsed.get('topic') or filename.rsplit('.', 1)[0]
        img_count = sum(1 for q in questions if q.get('img'))
        print(f"parse_document: {filename} -> {len(questions)} savol, {img_count} rasm", file=sys.stderr)
        return jsonify({'topic': topic, 'questions': questions, 'questionsToAsk': len(questions)})
    except Exception as e:
        print(f"parse_document error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


# ─── RASMDAN SAVOL O'QISH (Vision AI) ──────────────────────────────────────
# Test sahifasining fotosurati (.jpg/.png) tashlanganda - AI RASMNING OZINI
# korib, savol/variant/togri javobni aniqlaydi. Togri javob odatda nuqta,
# doira yoki boshqa belgi bilan korsatiladi - AI shu belgini izlaydi.
GROQ_VISION_MODEL = os.environ.get('GROQ_VISION_MODEL', 'qwen/qwen3.6-27b')

@app.route('/parse_image', methods=['POST', 'OPTIONS'])
def parse_image():
    if request.method == 'OPTIONS':
        return '', 200
    if not GROQ_API_KEY:
        return jsonify({'error': 'AI xizmati sozlanmagan (serverda GROQ_API_KEY yoq)'}), 500
    try:
        body = request.get_json(force=True, silent=True) or {}
        b64 = body.get('data', '')
        filename = body.get('filename', 'rasm')
        if not b64:
            return jsonify({'error': 'No data'}), 400
        mime = 'image/png' if filename.lower().endswith('.png') else 'image/jpeg'
        prompt_text = (
            "Bu rasmda test savoli(lari) va variantlari bor. Rasmni diqqat bilan tahlil qil.\n"
            "QATIY QOIDALAR:\n"
            "- Togri javob odatda variant yonida NUQTA, TOLDIRILGAN DOIRA, GALOCHKA (✓), "
            "AYLANTIRIB CHIZILGAN yoki shunga oxshash BELGI bilan korsatilgan boladi - shu "
            "belgi qoyilgan variantni TOGRI JAVOB deb bilgila.\n"
            "- Agar hech qanday variantda belgi topilmasa, oshu savolni OTKAZIB YUBOR - "
            "taxmin qilma.\n"
            "- Rasmda bir nechta savol bolsa, hammasini chiqar.\n"
            "- FAQAT JSON qaytar, boshqa hech narsa yozma:\n"
            "{\"questions\": [{\"text\": \"savol matni\", \"options\": [\"variant1\",\"variant2\"], "
            "\"correct\": [0], \"isMulti\": false}]}\n"
        )
        r = requests.post(
            'https://api.groq.com/openai/v1/chat/completions',
            headers={'Content-Type': 'application/json', 'Authorization': 'Bearer ' + GROQ_API_KEY},
            json={
                'model': GROQ_VISION_MODEL,
                'messages': [{
                    'role': 'user',
                    'content': [
                        {'type': 'text', 'text': prompt_text},
                        {'type': 'image_url', 'image_url': {'url': f'data:{mime};base64,{b64}'}},
                    ],
                }],
                'max_tokens': 4000,
                'temperature': 0.1,
                'response_format': {'type': 'json_object'},
            },
            timeout=60,
        )
        result = r.json()
        if 'error' in result:
            return jsonify({'error': f"AI xatosi (model: {GROQ_VISION_MODEL}): " + str(result['error'])}), 502
        content = result['choices'][0]['message']['content'].strip()
        if content.startswith('```'):
            content = re.sub(r'^```[a-zA-Z]*\n?', '', content)
            content = re.sub(r'```\s*$', '', content)
        parsed = json.loads(content)
        questions = parsed.get('questions', [])
        if not questions:
            return jsonify({'error': "Rasmdan savol aniqlanmadi (togri javob belgisi topilmagan bolishi mumkin)"}), 400
        for i, q in enumerate(questions):
            q['id'] = f'{int(time.time()*1000)}_{i}_{uuid.uuid4().hex[:6]}'
            q.setdefault('subject', 'math')
            if 'isMulti' not in q:
                q['isMulti'] = len(q.get('correct', [])) > 1
        topic = filename.rsplit('.', 1)[0]
        return jsonify({'topic': topic, 'questions': questions, 'questionsToAsk': len(questions)})
    except Exception as e:
        print(f"parse_image error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


# ─── KITOBDAN TEST TUZISH (AI generatsiya) ────────────────────────────────
# Bu yuqoridagi /parse_document'dan TUBDAN farq qiladi: u yerda AI mavjud
# savollarni MATNDAN TOPADI, bu yerda esa AI oddiy darslik/kitob matnidan
# YANGI savollarni OZI TUZADI. Kitob "mavzu"larga (boblarga) bolinadi -
# sarlavha uslubi (Word "Heading") yoki shrift olchami (PDF) orqali
# aniqlanadi; agar sarlavha topilmasa, teng bolimlarga bolinadi.

def extract_pdf_sections(data):
    """PDF'ni shrift olchami/qalinligi asosida mavzu(bob)larga ajratadi."""
    import fitz
    from collections import Counter
    doc = fitz.open(stream=data, filetype='pdf')
    all_lines = []
    for page in doc:
        blocks = page.get_text('dict').get('blocks', [])
        for b in blocks:
            if b.get('type') != 0:
                continue
            for line in b.get('lines', []):
                line_text, sizes, bolds = '', [], []
                for span in line.get('spans', []):
                    line_text += span.get('text', '')
                    sizes.append(span.get('size', 0))
                    bolds.append(bool(span.get('flags', 0) & 16))
                line_text = line_text.strip()
                if line_text:
                    avg_size = sum(sizes) / len(sizes) if sizes else 0
                    all_lines.append({'text': line_text, 'size': avg_size, 'bold': any(bolds)})
    doc.close()
    if not all_lines:
        return []
    size_counts = Counter(round(l['size']) for l in all_lines)
    body_size = size_counts.most_common(1)[0][0] if size_counts else 12

    sections, cur_title, cur_text = [], None, []
    for l in all_lines:
        is_heading = (l['size'] > body_size * 1.25 or (l['bold'] and l['size'] >= body_size)) \
            and len(l['text']) < 80 and len(l['text'].split()) < 12
        if is_heading:
            if cur_title and cur_text:
                sections.append({'title': cur_title, 'text': '\n'.join(cur_text)})
            cur_title, cur_text = l['text'], []
        else:
            cur_text.append(l['text'])
    if cur_title and cur_text:
        sections.append({'title': cur_title, 'text': '\n'.join(cur_text)})

    if not sections:
        full_text = '\n'.join(l['text'] for l in all_lines)
        sections = _split_into_word_chunks(full_text)
    sections = [s for s in sections if len(s['text'].split()) > 30]
    return sections


def extract_docx_sections(data):
    """DOCX'ni Word'ning "Heading" uslublari (yoki qalin qisqa qatorlar)
    asosida mavzu(bob)larga ajratadi."""
    import docx
    doc = docx.Document(io.BytesIO(data))
    sections, cur_title, cur_text = [], None, []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or '') if para.style else ''
        is_heading = style_name.lower().startswith('heading') or style_name.lower().startswith('title')
        if not is_heading and len(text) < 80 and len(text.split()) < 12:
            real_runs = [r for r in para.runs if r.text.strip()]
            if real_runs and all(r.bold for r in real_runs):
                is_heading = True
        if is_heading:
            if cur_title and cur_text:
                sections.append({'title': cur_title, 'text': '\n'.join(cur_text)})
            cur_title, cur_text = text, []
        else:
            cur_text.append(text)
    if cur_title and cur_text:
        sections.append({'title': cur_title, 'text': '\n'.join(cur_text)})

    if not sections:
        full_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        sections = _split_into_word_chunks(full_text)
    sections = [s for s in sections if len(s['text'].split()) > 30]
    return sections


def _split_into_word_chunks(full_text, chunk_words=2500):
    words = full_text.split()
    if not words:
        return []
    chunks = [' '.join(words[i:i+chunk_words]) for i in range(0, len(words), chunk_words)]
    return [{'title': f'Bo\'lim {i+1}', 'text': c} for i, c in enumerate(chunks)]


def generate_questions_from_content(text_chunk, topic_name, n_questions):
    if not GROQ_API_KEY:
        raise Exception('AI xizmati sozlanmagan (serverda GROQ_API_KEY yoq)')
    # MUHIM: Groq bepul tarifida llama-3.3-70b-versatile uchun 12000
    # token/daqiqa (TPM) limiti bor. Bu limit KIRISH (matn+prompt) VA
    # CHIQISH (max_tokens) yigindisiga taalluqli. Avval MAX_CHARS=12000 va
    # max_tokens=8000 birgalikda bu limitdan oshib ketardi (masalan 13053
    # token so'ralgan, 12000 ruxsat etilgan) - shuning uchun ikkalasi ham
    # kichraytirilgan.
    MAX_CHARS = 6000
    if len(text_chunk) > MAX_CHARS:
        text_chunk = text_chunk[:MAX_CHARS]
    n_questions = max(1, min(int(n_questions or 5), 20))
    prompt = (
        f"Sen tajribali ustozsan. Quyidagi darslik/oquv matni asosida ANIQ {n_questions} ta "
        "test savoli tuz.\n\n"
        "QATIY TALABLAR:\n"
        "- Savollar MURAKKAB va CHUQUR bolishi SHART - oddiy eslab qolish emas, balki "
        "tushunish, taqqoslash, tahlil qilish yoki qollash darajasida bolsin.\n"
        "- MUHIM - SAVOLLAR TURINI ARALASHTIR, bir xil qolipni takrorlama: "
        "\"... nima anglatadi?\" yoki \"... nima?\" kabi bitta qolipni HAMMA savolda ishlatish "
        "QATIYAN TAQIQLANADI. Buning ornida har xil savol turlaridan foydalan, masalan: "
        "sabab-natija (\"Nega ... sodir boldi?\"), qiyoslash (\"... bilan ... orasidagi asosiy farq "
        "nimada?\"), xronologiya/tartib (\"... dan keyin nima sodir boldi?\"), qollash "
        "(\"Agar ... bolsa, natija qanday bolar edi?\"), baholash (\"... ning eng muhim sababi "
        "nima?\"), togri/notogri tahlil (\"Quyidagilardan qaysi biri togri?\"). Har 5 ta savolning "
        "kamida 4 tasi BOSHQA-BOSHQA qolipda bolsin.\n"
        "- Har bir savolda ANIQ 4 ta variant, ulardan FAQAT BITTASI togri.\n"
        "- Notogri variantlar (distraktorlar) HAQIQATAN ishonarli va chalgituvchi bolsin - "
        "mavzuga oid, yuzaki qaraganda togriday tuyuladigan tushunchalar bolsin.\n"
        "- Savollar FAQAT berilgan matn mazmuniga asoslansin, undan tashqari malumot qoshma.\n"
        "- FAQAT JSON qaytar, boshqa hech narsa yozma:\n"
        "{\"questions\": [{\"text\": \"...\", \"options\": [\"...\",\"...\",\"...\",\"...\"], "
        "\"correct\": [0], \"isMulti\": false}]}\n\n"
        f"Mavzu: {topic_name}\n\n"
        "Matn:\n" + text_chunk
    )
    r = requests.post(
        'https://api.groq.com/openai/v1/chat/completions',
        headers={'Content-Type': 'application/json', 'Authorization': 'Bearer ' + GROQ_API_KEY},
        json={
            'model': 'llama-3.3-70b-versatile',
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': 4000,
            'temperature': 0.4,
            'response_format': {'type': 'json_object'},
        },
        timeout=90,
    )
    result = r.json()
    if 'error' in result:
        raise Exception(str(result['error']))
    content = result['choices'][0]['message']['content'].strip()
    if content.startswith('```'):
        content = re.sub(r'^```[a-zA-Z]*\n?', '', content)
        content = re.sub(r'```\s*$', '', content)
    parsed = json.loads(content)
    return parsed.get('questions', [])


def _run_generate_job(job_id, filename, data, mode, total_questions):
    try:
        ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        JOBS[job_id]['progress'] = "Fayldan matn va mavzular ajratilmoqda..."
        if ext == 'pdf':
            sections = extract_pdf_sections(data)
        elif ext == 'docx':
            sections = extract_docx_sections(data)
        elif ext == 'txt':
            sections = [{'title': filename.rsplit('.', 1)[0], 'text': data.decode('utf-8', errors='replace')}]
        else:
            JOBS[job_id]['status'] = 'error'
            JOBS[job_id]['error'] = 'Fayl turi qollab-quvvatlanmaydi (.pdf, .docx, .txt)'
            return
        if not sections:
            JOBS[job_id]['status'] = 'error'
            JOBS[job_id]['error'] = 'Fayldan yetarli matn topilmadi'
            return

        MAX_SECTIONS = 15
        if len(sections) > MAX_SECTIONS:
            sections = sections[:MAX_SECTIONS]

        results = []
        if mode == 'per_topic':
            for i, sec in enumerate(sections):
                JOBS[job_id]['progress'] = f'Savollar tuzilmoqda: "{sec["title"]}" ({i+1}/{len(sections)})...'
                qs = generate_questions_from_content(sec['text'], sec['title'], total_questions)
                if qs:
                    results.append({'topic': sec['title'], 'questions': qs, 'questionsToAsk': len(qs)})
        else:
            per_section = max(1, int(total_questions) // len(sections))
            all_qs = []
            for i, sec in enumerate(sections):
                JOBS[job_id]['progress'] = f'Savollar tuzilmoqda: bolim {i+1}/{len(sections)}...'
                qs = generate_questions_from_content(sec['text'], sec['title'], per_section)
                all_qs.extend(qs)
            topic_name = filename.rsplit('.', 1)[0]
            all_qs = all_qs[:int(total_questions)] if total_questions else all_qs
            results.append({'topic': topic_name, 'questions': all_qs, 'questionsToAsk': len(all_qs)})

        for res in results:
            for i, q in enumerate(res['questions']):
                q['id'] = f'{int(time.time()*1000)}_{i}_{uuid.uuid4().hex[:6]}'
                q.setdefault('subject', 'math')
                if 'isMulti' not in q:
                    q['isMulti'] = len(q.get('correct', [])) > 1

        JOBS[job_id]['status'] = 'done'
        JOBS[job_id]['results'] = results
        JOBS[job_id]['progress'] = 'Tugadi'
        total_qs = sum(len(r['questions']) for r in results)
        print(f"generate job {job_id}: {len(results)} mavzu, {total_qs} savol", file=sys.stderr)
    except Exception as e:
        print(f"generate job {job_id} error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        JOBS[job_id]['status'] = 'error'
        JOBS[job_id]['error'] = str(e)


@app.route('/generate_test_start', methods=['POST', 'OPTIONS'])
def generate_test_start():
    if request.method == 'OPTIONS':
        return '', 200
    body = request.get_json(force=True, silent=True) or {}
    b64 = body.get('data', '')
    filename = body.get('filename', 'kitob')
    mode = body.get('mode', 'overall')
    total_questions = int(body.get('total_questions', 20))
    if not b64:
        return jsonify({'error': 'No data'}), 400
    data = base64.b64decode(b64)
    job_id = str(uuid.uuid4())
    JOBS[job_id] = {'status': 'processing', 'progress': 'Boshlanmoqda...'}
    JOB_TIMESTAMPS[job_id] = _time.time()
    _cleanup_stale_jobs()
    t = threading.Thread(target=_run_generate_job, args=(job_id, filename, data, mode, total_questions), daemon=True)
    t.start()
    return jsonify({'job_id': job_id})


@app.route('/generate_test_status/<job_id>', methods=['GET'])
def generate_test_status(job_id):
    _cleanup_stale_jobs()
    job = JOBS.get(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    resp = {'status': job['status'], 'progress': job.get('progress', '')}
    if job['status'] == 'done':
        resp['results'] = job['results']
        JOBS.pop(job_id, None)
        JOB_TIMESTAMPS.pop(job_id, None)
    if job['status'] == 'error':
        resp['error'] = job.get('error', 'Nomalum xato')
        JOBS.pop(job_id, None)
        JOB_TIMESTAMPS.pop(job_id, None)
    return jsonify(resp)


@app.route('/parse_batch', methods=['POST', 'OPTIONS'])
def parse_batch():
    """Bir nechta faylni BIR SO'ROVDA qabul qiladi va BARCHA formulalarni
    faqat BITTA LibreOffice sessiyasida o'giradi. Eski (sinxron) versiya -
    Render'ning uzoq sorovlarni majburan uzib qoyishi sababli endi
    ishlatilmaydi, lekin orqaga moslik uchun qoldirilgan."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        body = request.get_json(force=True, silent=True) or {}
        files = body.get('files', [])
        if not files:
            return jsonify({'error': 'No files'}), 400
        file_results, file_emf_tasks = _process_files_raw(files)
        _resolve_batch_emfs(file_results, file_emf_tasks)
        return jsonify({'results': file_results})
    except Exception as e:
        print(f"Batch error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


def _process_files_raw(files):
    file_results = []
    file_emf_tasks = []
    for f in files:
        fname = f.get('filename', 'fayl')
        try:
            data = base64.b64decode(f.get('data', ''))
            xml_text = find_xml(data)
            if not xml_text:
                file_results.append({'filename': fname, 'error': 'XML not found'})
                file_emf_tasks.append([])
                continue
            topic, questions, emf_tasks, qta = extract_questions_raw(xml_text, data)
            file_results.append({'filename': fname, 'topic': topic, 'questions': questions, 'questionsToAsk': qta})
            file_emf_tasks.append(emf_tasks)
        except Exception as e:
            print(f"  {fname}: xato {e}", file=sys.stderr)
            file_results.append({'filename': fname, 'error': str(e)})
            file_emf_tasks.append([])
    return file_results, file_emf_tasks


def _resolve_batch_emfs(file_results, file_emf_tasks, job_id=None):
    global_list = []
    global_map = []
    for fi, tasks in enumerate(file_emf_tasks):
        for t in tasks:
            global_list.append((len(global_list), t['emf']))
            global_map.append((fi, t))
    total = len(global_list)
    if not total:
        return
    if job_id:
        JOBS[job_id]['progress'] = f'0/{total} rasm/formula ogirilmoqda...'

    # Kattaroq royxatlarni kichikroq boliklarga bolib ishlaymiz - shunda
    # foydalanuvchi progressni real vaqtda kora oladi (qotib qolganday
    # tuyulmasligi uchun), va bitta LibreOffice chaqiruvi haddan tashqari
    # katta bolib ketmaydi.
    # MUHIM: Render "Free" tarifi (512 MB RAM, 0.1 vCPU) uchun bu qiymat
    # ANIQ kamaytirilgan - avval 90 edi, endi 30. Kattaroq CHUNK bir vaqtning
    # ozida juda kop EMF'ni xotiraga yuklab, OOM (xotira tugashi) sabab
    # butun serverni qulatib qoyishi mumkin edi.
    CHUNK = 30
    for start in range(0, total, CHUNK):
        chunk = global_list[start:start+CHUNK]
        chunk_results = convert_all_emfs(chunk)
        for gidx, b64 in chunk_results.items():
            fi, t = global_map[gidx]
            res = file_results[fi]
            if 'questions' not in res:
                continue
            if t['kind'] == 'q':
                res['questions'][t['q_idx']]['image'] = b64
            else:
                opts = res['questions'][t['q_idx']]['options']
                if t['opt_idx'] < len(opts):
                    opts[t['opt_idx']] = b64
        done = min(start+CHUNK, total)
        if job_id:
            JOBS[job_id]['progress'] = f'{done}/{total} rasm/formula ogirildi...'
        print(f"  EMF progress: {done}/{total}", file=sys.stderr)
    return


# ─── FON VAZIFA (BACKGROUND JOB) TIZIMI ────────────────────────────────────
# Render.com (va boshqa hosting'lar) uzoq davom etadigan HTTP sorovlarni
# ozi majburan uzib qoyadi (odatda 30-100 soniyadan keyin), garchi bizning
# kod hali ishlab turgan bolsa ham. Buni chetlab otish uchun: katta ishni
# ORQA FONDA (alohida thread'da) qilamiz, brauzer esa tez-tez "tayyor
# bo'ldimi?" deb sorab turadi (polling). Har bir sorovning ozi tez
# (sub-sekund) bolgani uchun Render uni hech qachon uzib qoymaydi.
import threading, uuid

JOBS = {}  # job_id -> {'status':'processing'|'done'|'error', 'progress':str, 'results':[...], 'error':str}

def _run_batch_job(job_id, files):
    try:
        JOBS[job_id]['progress'] = f'0/{len(files)} fayl oqildi'
        file_results = []
        file_emf_tasks = []
        for i, f in enumerate(files):
            fname = f.get('filename', 'fayl')
            try:
                data = base64.b64decode(f.get('data', ''))
                xml_text = find_xml(data)
                if not xml_text:
                    file_results.append({'filename': fname, 'error': 'XML not found'})
                    file_emf_tasks.append([])
                else:
                    topic, questions, emf_tasks, qta = extract_questions_raw(xml_text, data)
                    file_results.append({'filename': fname, 'topic': topic, 'questions': questions, 'questionsToAsk': qta})
                    file_emf_tasks.append(emf_tasks)
            except Exception as e:
                print(f"  {fname}: xato {e}", file=sys.stderr)
                file_results.append({'filename': fname, 'error': str(e)})
                file_emf_tasks.append([])
            JOBS[job_id]['progress'] = f'{i+1}/{len(files)} fayl oqildi'

        _resolve_batch_emfs(file_results, file_emf_tasks, job_id)

        JOBS[job_id]['status'] = 'done'
        JOBS[job_id]['results'] = file_results
        JOBS[job_id]['progress'] = 'Tugadi'
        ok = sum(1 for r in file_results if 'questions' in r)
        print(f"Job {job_id}: tugadi, {ok}/{len(files)} muvaffaqiyatli", file=sys.stderr)
    except Exception as e:
        print(f"Job {job_id} error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        JOBS[job_id]['status'] = 'error'
        JOBS[job_id]['error'] = str(e)


# ─── TAQSIMLANGAN IMPORT (server-pool, umumiy EMF navbati) ────────────────
# ESKI USUL: har bir fayl-guruhi BITTA belgilangan serverga to'liq (XML
# ajratish + EMF konvertatsiya) yuborilar edi - shuning uchun ko'p kichik
# fayl bolganda ular navbat-navbat, faqat bitta guruh miqdorida parallel
# ishlanardi.
#
# YANGI USUL (2026-08): ikki bosqichga bolinadi -
#   1) EXTRACT (tez, faqat regex, LibreOffice YOQ) - istalgan BITTA serverda
#      barcha fayllar XML'dan ochiladi, EMF baytlari (hali PNG'ga
#      aylanmagan xom EMF) global royxatga chiqariladi.
#   2) CONVERT (sekin, LibreOffice) - shu global royxat (bitta "umumiy
#      guruh") frontend tomonidan BARCHA ROYXATDAGI SERVERLARGA
#      (Oracle VM1-6, Koyeb va h.k.) bir vaqtda, har birining tezligiga
#      qarab bolib-bolib yuboriladi ("work-stealing"). Har bir server
#      bu yerdagi /convert_emfs orqali ozining ulushini konvertatsiya
#      qiladi - qaysi faylga tegishli ekanligini BILISHI SHART EMAS,
#      faqat xom EMF baytini PNG'ga aylantiradi. Shu sabab bu endpoint
#      har qanday server turida (hatto zaif Koyeb'da ham) ishlay oladi.
def _b64_extract_emf_tasks(file_results, file_emf_tasks):
    """file_emf_tasks (python obyektlar, xom bayt) ni JSON-safe (base64)
    korinishga otkazadi, hamda har bir EMF vazifasiga GLOBAL noyob id
    beradi - frontend keyinchalik shu id orqali natijani togri joyga
    (togri fayl/savol/variantga) qaytarib qoya oladi."""
    out_tasks = []
    for fi, tasks in enumerate(file_emf_tasks):
        for t in tasks:
            out_tasks.append({
                'id': f'{fi}_{t["kind"]}_{t["q_idx"]}_{t.get("opt_idx", "")}',
                'file_index': fi,
                'emf_b64': base64.b64encode(t['emf']).decode(),
            })
    return out_tasks


@app.route('/extract_batch_start', methods=['POST', 'OPTIONS'])
def extract_batch_start():
    """1-bosqich: fayllarni XML'dan ochadi, EMF baytlarini (KONVERTATSIYA
    QILMASDAN) global royxat sifatida qaytaradi. Tez ishlaydi (regex),
    shuning uchun job/polling shart emas - sinxron javob beradi."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        body = request.get_json(force=True, silent=True) or {}
        files = body.get('files', [])
        if not files:
            return jsonify({'error': 'No files'}), 400
        file_results, file_emf_tasks = _process_files_raw(files)
        emf_tasks = _b64_extract_emf_tasks(file_results, file_emf_tasks)
        # Xotira tejash uchun - endi frontend'ga xom EMF baytlarini emas,
        # faqat FAYL natijalarini (savol matni, variantlar, hali rasmsiz
        # "bosh joy"lar bilan) va alohida emf_tasks royxatini qaytaramiz.
        return jsonify({'files': file_results, 'emf_tasks': emf_tasks})
    except Exception as e:
        print(f"extract_batch_start error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


@app.route('/convert_emfs', methods=['POST', 'OPTIONS'])
def convert_emfs():
    """2-bosqich: UNIVERSAL worker. Qaysi fayl/savolga tegishli ekanini
    bilmaydi - shunchaki {id, emf_b64} royxatini oladi, {id: png_data_url}
    qaytaradi. Istalgan server (kuchli yoki zaif) shu endpointni chaqirib,
    ozi ko'tara oladigan miqdorda ish olishi mumkin - shuning uchun
    frontend BITTA umumiy EMF navbatini kop turdagi serverlar orasida
    ularning tezligiga qarab bolib bera oladi."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        body = request.get_json(force=True, silent=True) or {}
        tasks = body.get('tasks', [])
        if not tasks:
            return jsonify({'error': 'No tasks'}), 400
        # MUHIM: bitta sorovda haddan tashqari kop EMF kelib, zaif server
        # (masalan Koyeb, 512MB RAM) OOM bolib qolmasligi uchun qattiq
        # yuqori chegara qoyilgan - frontend har doim kichikroq chunk
        # yuborishi kerak, lekin server tomonda ham ikkinchi himoya qatlami.
        MAX_TASKS_PER_REQUEST = 40
        if len(tasks) > MAX_TASKS_PER_REQUEST:
            return jsonify({'error': f'Juda kop EMF (max {MAX_TASKS_PER_REQUEST})'}), 400
        emf_list = []
        id_map = {}
        for idx, t in enumerate(tasks):
            try:
                emf_bytes = base64.b64decode(t['emf_b64'])
            except Exception:
                continue
            emf_list.append((idx, emf_bytes))
            id_map[idx] = t['id']
        results_by_idx = convert_all_emfs(emf_list)
        results = {id_map[idx]: b64 for idx, b64 in results_by_idx.items() if idx in id_map}
        return jsonify({'results': results})
    except Exception as e:
        print(f"convert_emfs error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': str(e)}), 500


@app.route('/parse_batch_start', methods=['POST', 'OPTIONS'])
def parse_batch_start():
    if request.method == 'OPTIONS':
        return '', 200
    body = request.get_json(force=True, silent=True) or {}
    files = body.get('files', [])
    if not files:
        return jsonify({'error': 'No files'}), 400
    job_id = str(uuid.uuid4())
    JOBS[job_id] = {'status': 'processing', 'progress': 'Boshlanmoqda...'}
    JOB_TIMESTAMPS[job_id] = _time.time()
    _cleanup_stale_jobs()
    t = threading.Thread(target=_run_batch_job, args=(job_id, files), daemon=True)
    t.start()
    return jsonify({'job_id': job_id})


@app.route('/parse_batch_status/<job_id>', methods=['GET'])
def parse_batch_status(job_id):
    _cleanup_stale_jobs()
    job = JOBS.get(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    resp = {'status': job['status'], 'progress': job.get('progress', '')}
    if job['status'] == 'done':
        resp['results'] = job['results']
        # Natija olib bolindi - xotirani bosh qilish uchun jobni ochiramiz.
        # (Rasmlar bilan togla katta bolgani uchun, xotirada qoldirib
        # qoyish serverni "toldirib" qoyishi mumkin edi.)
        JOBS.pop(job_id, None)
        JOB_TIMESTAMPS.pop(job_id, None)
    if job['status'] == 'error':
        resp['error'] = job.get('error', 'Nomalum xato')
        JOBS.pop(job_id, None)
        JOB_TIMESTAMPS.pop(job_id, None)
    return jsonify(resp)

# Xavfsizlik uchun: agar biror sababdan mijoz natijani hech qachon
# so'ramasa (masalan brauzer yopilib qolsa), 2 soatdan keyin eski
# joblarni avtomatik tozalaymiz - xotira sekin-asta toldirilmasin.
import time as _time
JOB_TIMESTAMPS = {}
def _cleanup_stale_jobs():
    now = _time.time()
    stale = [jid for jid, ts in list(JOB_TIMESTAMPS.items()) if now - ts > 7200]
    for jid in stale:
        JOBS.pop(jid, None)
        JOB_TIMESTAMPS.pop(jid, None)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))


# ─── OTA-ONA TELEGRAM BOTI ──────────────────────────────────────────────────
# Maqsad: ota-ona hech qanday login/parolsiz, bitta Telegram havolasini
# bosish orqali bolasining haftalik progressini avtomatik oladi. Token
# HECH QACHON kodda yozilmaydi - faqat Render Environment'da saqlanadi
# (Dashboard > Environment > TELEGRAM_BOT_TOKEN).
TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', '')

# ─── EMAIL (2026-08-29: ustoz royxatdan otish/tasdiqlash xabarlari uchun) ──
# Gmail SMTP orqali ishlaydi (App Password bilan - oddiy Gmail parol EMAS,
# Google hisobida "App passwords" bolimidan 16 xonali maxsus parol olinadi).
# Kalitlar Render Environment'da: EMAIL_SMTP_USER (Gmail manzil),
# EMAIL_SMTP_PASS (16 xonali App Password).
EMAIL_SMTP_USER = os.environ.get('EMAIL_SMTP_USER', '')
EMAIL_SMTP_PASS = os.environ.get('EMAIL_SMTP_PASS', '')

def _send_email(to_email, subject, body_html):
    if not (EMAIL_SMTP_USER and EMAIL_SMTP_PASS):
        print("Email: EMAIL_SMTP_USER/EMAIL_SMTP_PASS yoq, xabar yuborilmadi", file=sys.stderr)
        return False
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = f'EduTest Pro <{EMAIL_SMTP_USER}>'
        msg['To'] = to_email
        msg.attach(MIMEText(body_html, 'html'))
        with smtplib.SMTP_SSL('smtp.gmail.com', 465, timeout=15) as server:
            server.login(EMAIL_SMTP_USER, EMAIL_SMTP_PASS)
            server.sendmail(EMAIL_SMTP_USER, [to_email], msg.as_string())
        return True
    except Exception as e:
        print(f"Email yuborishda xato: {e}", file=sys.stderr)
        return False

@app.route('/notify_teacher_status', methods=['POST', 'OPTIONS'])
def notify_teacher_status():
    """Admin ustozni tasdiqlagan yoki rad etganda chaqiriladi - ustozning
    ro'yxatdan o'tishda ko'rsatgan emailiga avtomatik xabar yuboradi."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        body = request.get_json(force=True, silent=True) or {}
        to_email = (body.get('email') or '').strip()
        name = (body.get('name') or '').strip()
        login = (body.get('login') or '').strip()
        status = (body.get('status') or '').strip()
        if not to_email or status not in ('approved', 'blocked'):
            return jsonify({'error': 'email va togri status (approved/blocked) kerak'}), 400
        if status == 'approved':
            subject = 'EduTest Pro - Hisobingiz tasdiqlandi'
            body_html = (
                f"<p>Assalomu alaykum, <b>{name}</b>!</p>"
                f"<p>Sizning EduTest Pro'dagi so'rovingiz <b>tasdiqlandi</b>. Endi tizimga kirishingiz mumkin.</p>"
                f"<p>Login: <b>{login}</b></p>"
                f"<p>Ro'yxatdan o'tishda kiritgan parolingiz bilan kiring.</p>"
            )
        else:
            subject = 'EduTest Pro - So\'rovingiz haqida'
            body_html = (
                f"<p>Assalomu alaykum, <b>{name}</b>!</p>"
                f"<p>Afsuski, sizning EduTest Pro'dagi ustoz sifatida ro'yxatdan o'tish so'rovingiz "
                f"admin tomonidan <b>rad etildi</b>.</p>"
                f"<p>Savollaringiz bo'lsa, maktab administratsiyasi bilan bog'laning.</p>"
            )
        ok = _send_email(to_email, subject, body_html)
        return jsonify({'sent': ok})
    except Exception as e:
        print(f"notify_teacher_status xato: {e}", file=sys.stderr)
        return jsonify({'error': str(e)}), 500
TELEGRAM_BOT_USERNAME = os.environ.get('TELEGRAM_BOT_USERNAME', 'EduTestProrobot')
# Haftalik yuborishni tashqi cron (masalan cron-job.org) chaqiradi -
# Render'ning bepul tarifida doimiy ishlaydigan fon jarayon (APScheduler)
# ishonchli emas (server uxlab qolishi mumkin), shuning uchun tashqi
# "uyg'otuvchi" so'rov + maxfiy kalit orqali himoyalangan endpoint ishlatiladi.
TG_CRON_SECRET = os.environ.get('TG_CRON_SECRET', '')

def _sb_headers():
    return {
        'apikey': SUPABASE_SERVICE_KEY,
        'Authorization': 'Bearer ' + SUPABASE_SERVICE_KEY,
        'Content-Type': 'application/json',
    }

def _load_main_db():
    """edutest_store ichidagi asosiy JSON blobni service key bilan o'qiydi -
    db_get bilan bir xil, lekin ichki ishlatish uchun to'g'ridan-to'g'ri
    dict qaytaradi (JSON string emas)."""
    r = requests.get(
        f'{SUPABASE_URL}/rest/v1/edutest_store',
        params={'key': f'eq.{DB_KEY}', 'select': 'value'},
        headers=_sb_headers(), timeout=15
    )
    r.raise_for_status()
    data = r.json()
    if not data:
        return {}
    return json.loads(data[0]['value'])

def _tg_send(chat_id, text, keyboard=None):
    if not TELEGRAM_BOT_TOKEN:
        print("Telegram: TELEGRAM_BOT_TOKEN yoq, xabar yuborilmadi", file=sys.stderr)
        return False
    try:
        payload = {'chat_id': chat_id, 'text': text, 'parse_mode': 'HTML'}
        if keyboard:
            payload['reply_markup'] = {'inline_keyboard': keyboard}
        r = requests.post(
            f'https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage',
            json=payload, timeout=15
        )
        return r.status_code == 200
    except Exception as e:
        print(f"Telegram send xato: {e}", file=sys.stderr)
        return False

def _tg_answer_callback(callback_id, text=''):
    if not TELEGRAM_BOT_TOKEN:
        return
    try:
        requests.post(
            f'https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/answerCallbackQuery',
            json={'callback_query_id': callback_id, 'text': text}, timeout=10
        )
    except Exception:
        pass

def _tg_get_links_for_chat(chat_id):
    """Shu Telegram chat_id'ga ulangan barcha o'quvchilarni qaytaradi (bir
    ota-onaning bir nechta farzandi bolishi mumkin, hammasi shu botga
    ulangan bolishi mumkin)."""
    r = requests.get(
        f'{SUPABASE_URL}/rest/v1/parent_links',
        params={'telegram_chat_id': f'eq.{chat_id}', 'select': '*'},
        headers=_sb_headers(), timeout=15
    )
    return r.json() if r.status_code < 400 else []

def _student_name(db, login):
    student = next((u for u in db.get('users', []) if u.get('login') == login), None)
    return student['name'] if student else login

def _tg_send_menu(chat_id, greeting=None):
    """Asosiy menyu - ulangan har bir bola uchun tugma + umumiy amallar.
    Har qanday tushunarsiz xabar/komanda kelganda ham shu menyu qayta
    ko'rsatiladi - shunda ota-ona botni 'ishlamayapti' deb blocklamaydi,
    doim biror tugma bosib ishlatishi mumkin."""
    links = _tg_get_links_for_chat(chat_id)
    if not links:
        _tg_send(chat_id, (greeting or '') + "\n\nHozircha hech qanday farzand ulanmagan. Ulash uchun EduTest Pro ilovasida farzandingizning profilidan 'Ota-onani ulash' havolasini oching.")
        return
    db = _load_main_db()
    kb = []
    for link in links:
        sname = _student_name(db, link['student_login'])
        kb.append([{'text': f"📊 {sname} — hozirgi holat", 'callback_data': f"progress:{link['student_login']}"}])
    kb.append([{'text': "❓ Yordam", 'callback_data': "help"}])
    kb.append([{'text': "🔕 Xabarlarni to'xtatish", 'callback_data': "unlink_menu"}])
    text = (greeting + "\n\n" if greeting else "") + "Quyidagi tugmalardan birini tanlang:"
    _tg_send(chat_id, text, keyboard=kb)

def _tg_send_help(chat_id):
    _tg_send(chat_id, (
        "ℹ️ <b>EduTest Pro ota-ona boti</b>\n\n"
        "Bu bot farzandingizning maktabdagi test natijalari haqida avtomatik xabar beradi:\n"
        "• Har yakshanba kechqurun — haftalik hisobot avtomatik keladi\n"
        "• Istalgan vaqt \"Hozirgi holat\" tugmasini bosib, so'nggi 7 kunlik natijani darhol ko'rishingiz mumkin\n\n"
        "Botni to'xtatish uchun \"Xabarlarni to'xtatish\" tugmasidan foydalaning yoki botni bloklang.\n\n"
        "Savollaringiz bo'lsa, farzandingizning ustoziga murojaat qiling — bot savollarga javob bermaydi, faqat avtomatik hisobot yuboradi."
    ))
    _tg_send_menu(chat_id)

@app.route('/parent_link_create', methods=['POST', 'OPTIONS'])
def parent_link_create():
    """Ustoz/o'quvchi 'Ota-onani ulash' tugmasini bosganda chaqiriladi.
    Bitta bir martalik kod yaratadi va Telegram deep-link qaytaradi -
    ota-ona shu havolani bosishi bilan bot avtomatik bog'laydi."""
    if request.method == 'OPTIONS':
        return '', 200
    if not SUPABASE_SERVICE_KEY:
        return jsonify({'error': 'Server sozlanmagan (SUPABASE_SERVICE_KEY yoq)'}), 500
    try:
        body = request.get_json(force=True, silent=True) or {}
        student_login = (body.get('student_login') or '').strip()
        if not student_login:
            return jsonify({'error': 'student_login kerak'}), 400
        code = uuid.uuid4().hex[:10]
        r = requests.post(
            f'{SUPABASE_URL}/rest/v1/parent_links',
            headers=_sb_headers(),
            json={'student_login': student_login, 'link_code': code},
            timeout=15
        )
        if r.status_code >= 400:
            return jsonify({'error': f'Supabase xato: {r.status_code} {r.text[:200]}'}), 502
        link = f'https://t.me/{TELEGRAM_BOT_USERNAME}?start={code}'
        return jsonify({'link': link, 'code': code})
    except Exception as e:
        print(f"parent_link_create xato: {e}", file=sys.stderr)
        return jsonify({'error': str(e)}), 500

@app.route('/tg_webhook', methods=['POST'])
def tg_webhook():
    """Telegram'dan kelgan har bir update shu yerga tushadi (webhook orqali,
    polling emas - Render'da doimiy ishlaydigan jarayon shart emas)."""
    try:
        update = request.get_json(force=True, silent=True) or {}

        # ── TUGMA BOSILGANDA (callback_query) ──────────────────────────
        cq = update.get('callback_query')
        if cq:
            chat_id = (cq.get('message') or {}).get('chat', {}).get('id')
            data = cq.get('data') or ''
            _tg_answer_callback(cq.get('id', ''))
            if not chat_id:
                return jsonify({'ok': True})

            if data == 'help':
                _tg_send_help(chat_id)
            elif data == 'unlink_menu':
                links = _tg_get_links_for_chat(chat_id)
                if not links:
                    _tg_send(chat_id, "Ulangan farzand topilmadi.")
                else:
                    db = _load_main_db()
                    kb = [[{'text': f"❌ {_student_name(db, l['student_login'])}ni uzish", 'callback_data': f"unlink:{l['student_login']}"}] for l in links]
                    kb.append([{'text': "⬅️ Orqaga", 'callback_data': "menu"}])
                    _tg_send(chat_id, "Qaysi farzand uchun xabarlarni to'xtatmoqchisiz?", keyboard=kb)
            elif data.startswith('unlink:'):
                login = data.split(':', 1)[1]
                links = _tg_get_links_for_chat(chat_id)
                target = next((l for l in links if l['student_login'] == login), None)
                if target:
                    requests.patch(
                        f'{SUPABASE_URL}/rest/v1/parent_links',
                        params={'id': f'eq.{target["id"]}'},
                        headers=_sb_headers(),
                        json={'telegram_chat_id': None},
                        timeout=15
                    )
                db = _load_main_db()
                _tg_send(chat_id, f"✅ {_student_name(db, login)} uchun xabarlar to'xtatildi. Qayta ulanish uchun ilovadan yangi havola oling.")
                _tg_send_menu(chat_id)
            elif data == 'menu':
                _tg_send_menu(chat_id)
            elif data.startswith('progress:'):
                login = data.split(':', 1)[1]
                db = _load_main_db()
                summary = _weekly_summary_for_student(db, login, on_demand=True)
                _tg_send(chat_id, summary or "Ma'lumot topilmadi.")
                _tg_send_menu(chat_id)
            return jsonify({'ok': True})

        # ── ODDIY XABAR/KOMANDA ─────────────────────────────────────────
        msg = update.get('message') or {}
        text = (msg.get('text') or '').strip()
        chat_id = (msg.get('chat') or {}).get('id')
        if not chat_id:
            return jsonify({'ok': True})

        if text.startswith('/start'):
            parts = text.split(maxsplit=1)
            code = parts[1].strip() if len(parts) > 1 else ''
            if not code:
                # Kod bermasdan /start bossa - agar u avval ulangan bolsa,
                # to'g'ridan-to'g'ri menyuni ko'rsatamiz (yangi ulash shart
                # emasligini tushunsin, bot "ishlamayapti" deb otylamasin).
                existing = _tg_get_links_for_chat(chat_id)
                if existing:
                    _tg_send_menu(chat_id, greeting="Assalomu alaykum! Qaytib kelganingiz uchun rahmat.")
                else:
                    _tg_send(chat_id, "Assalomu alaykum! Bu — EduTest Pro ota-ona boti.\n\nBolangizning o'qituvchisidan yoki EduTest Pro ilovasidan olingan shaxsiy havola orqali ulaning.")
                return jsonify({'ok': True})
            # Kodni parent_links jadvalida topib, shu chatga bog'laymiz
            r = requests.get(
                f'{SUPABASE_URL}/rest/v1/parent_links',
                params={'link_code': f'eq.{code}', 'select': '*'},
                headers=_sb_headers(), timeout=15
            )
            rows = r.json() if r.status_code < 400 else []
            if not rows:
                _tg_send(chat_id, "Kechirasiz, havola topilmadi yoki eskirgan. Iltimos, ilovadan yangi havola so'rang.")
                return jsonify({'ok': True})
            row = rows[0]
            requests.patch(
                f'{SUPABASE_URL}/rest/v1/parent_links',
                params={'id': f'eq.{row["id"]}'},
                headers=_sb_headers(),
                json={'telegram_chat_id': str(chat_id), 'linked_at': 'now()'},
                timeout=15
            )
            db = _load_main_db()
            sname = _student_name(db, row['student_login'])
            _tg_send_menu(chat_id, greeting=f"✅ Muvaffaqiyatli ulandi! <b>{sname}</b>ning haftalik progressi endi shu yerga avtomatik yuborib turiladi.")
            return jsonify({'ok': True})

        if text.startswith('/help'):
            _tg_send_help(chat_id)
            return jsonify({'ok': True})
        if text.startswith('/menu'):
            _tg_send_menu(chat_id)
            return jsonify({'ok': True})

        # Har qanday boshqa (tushunarsiz) xabar - hech qachon "faqat
        # avtomatik xabar" deb qisqa javob bermaymiz endi, chunki bu ota-
        # onaga bot "o'lik/ishlamayapti" tuyuladi. Buning o'rniga har doim
        # ishlatsa bo'ladigan MENYUNI ko'rsatamiz.
        _tg_send_menu(chat_id)
        return jsonify({'ok': True})
    except Exception as e:
        print(f"tg_webhook xato: {e}", file=sys.stderr)
        return jsonify({'ok': True})  # Telegram'ga har doim 200 qaytaramiz, aks holda qayta-qayta urinaveradi

def _weekly_summary_for_student(db, login, on_demand=False):
    """Bitta o'quvchi uchun so'nggi 7 kunlik natijalar asosida qisqa,
    ota-ona uchun tushunarli (raqamlarsiz emas, lekin sodda) xabar tuzadi.
    on_demand=True bolsa (ota-ona 'Hozirgi holat' tugmasini bossa) matn
    sarlavhasi 'haftalik hisobot' emas, 'joriy holat' deb chiqadi - chunki
    bu haftalik avtomatik xabar emas, ota-ona o'zi so'rab olgan."""
    import datetime
    now = datetime.datetime.now()
    cutoff = now - datetime.timedelta(days=7)
    student = next((u for u in db.get('users', []) if u.get('login') == login), None)
    if not student:
        return None
    name = student.get('name', login)
    results = [r for r in db.get('results', []) if r.get('student') == name and r.get('mode') == 'exam']

    def parse_uz_date(s):
        try:
            d, m, y = s.split('.')
            return datetime.datetime(int(y), int(m), int(d))
        except Exception:
            return None

    recent = [r for r in results if (parse_uz_date(r.get('date', '')) or now) >= cutoff]
    header = f"📊 <b>{name}</b>ning joriy holati (so'nggi 7 kun):\n" if on_demand else f"📊 <b>{name}</b>ning bu haftalik hisoboti:\n"
    if not recent:
        return header + f"\nSo'nggi 7 kunda <b>{name}</b> hech qanday imtihon topshirmadi. Davomiylik natijalarni yaxshilashning eng muhim omili — bir necha daqiqa ajratishni tavsiya qilamiz."

    topics = sorted(set(r['topic'] for r in recent))
    avg = round(sum(r['pct'] for r in recent) / len(recent))
    weakest = min(recent, key=lambda r: r['pct'])
    lines = [header]
    lines.append(f"✅ {len(recent)} ta imtihon topshirdi, {len(topics)} ta mavzu bo'yicha")
    lines.append(f"📈 O'rtacha natija: <b>{avg}%</b>")
    if weakest['pct'] < 60:
        lines.append(f"⚠️ Eng qiyin joyi: <b>{weakest['topic']}</b> ({weakest['pct']}%) — shu mavzuni birga takrorlashni tavsiya qilamiz.")
    else:
        lines.append("👏 Barcha mavzularda yaxshi natija!")
    return '\n'.join(lines)

@app.route('/tg_send_weekly', methods=['POST', 'GET'])
def tg_send_weekly():
    """Tashqi cron (masalan cron-job.org, haftada 1 marta - yakshanba
    kechqurun) shu endpointni maxfiy kalit bilan chaqiradi. Bu Render'da
    doimiy ishlab turadigan scheduler shart emasligini anglatadi."""
    secret = request.args.get('secret', '')
    if not TG_CRON_SECRET or secret != TG_CRON_SECRET:
        return jsonify({'error': 'Ruxsat yoq'}), 403
    if not SUPABASE_SERVICE_KEY or not TELEGRAM_BOT_TOKEN:
        return jsonify({'error': 'Server sozlanmagan'}), 500
    try:
        db = _load_main_db()
        r = requests.get(
            f'{SUPABASE_URL}/rest/v1/parent_links',
            params={'telegram_chat_id': 'not.is.null', 'select': '*'},
            headers=_sb_headers(), timeout=15
        )
        links = r.json() if r.status_code < 400 else []
        sent, failed = 0, 0
        seen_pairs = set()
        for link in links:
            key = (link['student_login'], link['telegram_chat_id'])
            if key in seen_pairs:
                continue
            seen_pairs.add(key)
            summary = _weekly_summary_for_student(db, link['student_login'])
            if not summary:
                continue
            if _tg_send(link['telegram_chat_id'], summary):
                sent += 1
            else:
                failed += 1
        return jsonify({'sent': sent, 'failed': failed, 'total_links': len(seen_pairs)})
    except Exception as e:
        print(f"tg_send_weekly xato: {e}", file=sys.stderr)
        return jsonify({'error': str(e)}), 500

